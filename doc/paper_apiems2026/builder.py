#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""APIEMS FullPaperTemplate.docx へ原稿を流し込む docx ビルダー。

テンプレートの書式を実測した以下の仕様を再現する:
  - タイトル: 中央揃え TNR 20pt / 著者: 中央揃え TNR 10pt bold / 所属: 非bold
  - Abstract: スタイル"2"(heading2) + bold解除 + 左右インデント680twips, 10pt TNR
  - 見出し: スタイル"2" + TNR 11pt (sz=22)
  - 本文: スタイル"a8"(Body Text Indent, TNR 10pt, 1字下げ) + 行送り 240 exact
  - 1段(タイトル部)→2段(本文) は continuous sectPr、全幅図は 1段 continuous 区間
  - 数式: OMML(インライン) + タブで式番号右寄せ

コンテンツはブロックのリストで与える（content_ja.py / content_en.py 参照）:
  ('title', str) ('authors', [..]) ('affil', [..]) ('abstract', str)
  ('keywords', str) ('h1'|'h2'|'h3', str) ('p', str) ('p_noindent', str)
  ('eq', omml_xml_str, '(1)') ('fig', png_path, caption_str, 'full'|'col', height_scale)
  ('table', header_row, rows, col_widths_ratio) ('refs_heading', str) ('ref', str)
インライン記法: **bold** / *italic* / $math$（_ ^ 添字, 英字イタリック）
"""
import copy
import os
import re

from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import qn, nsmap
from docx.shared import Emu, Twips

NS_W = 'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
NS_M = 'xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"'

# 段組寸法（テンプレ実測, twips）
PAGE_W, MAR_L, MAR_R, COL_GAP = 11906, 964, 567, 425
TEXT_W = PAGE_W - MAR_L - MAR_R              # 10375
COL_W = (TEXT_W - COL_GAP) // 2              # 4975

SECT_PGDIMS = ('<w:pgSz w:w="11906" w:h="16838" w:code="9"/>'
               '<w:pgMar w:top="1701" w:right="567" w:bottom="2268" w:left="964"'
               ' w:header="851" w:footer="1134" w:gutter="0"/>')
SECT_GRID = '<w:docGrid w:type="linesAndChars" w:linePitch="360"/>'


def sect_xml(cols):
    c = ('<w:cols w:num="2" w:space="425"/>' if cols == 2
         else '<w:cols w:space="720"/>')
    return (f'<w:sectPr {NS_W}><w:type w:val="continuous"/>'
            f'{SECT_PGDIMS}{c}{SECT_GRID}</w:sectPr>')


# ---------------- インライン記法 → runs ----------------

_TOKEN = re.compile(r'(\*\*.+?\*\*|\*[^*]+?\*|\$[^$]+?\$)', re.S)

# バックスラッシュでエスケープした \* \$ はインライン記法として解釈させず
# 文字そのものとして出力する（有意水準凡例 *p<.05 等のため）。
# 一旦センチネル(私用領域)へ退避し、_esc の最終段で復元する。
_ESC_STAR, _ESC_DOLLAR = '', ''


def _protect_escapes(t):
    return t.replace(r'\*', _ESC_STAR).replace(r'\$', _ESC_DOLLAR)


def _esc(t):
    return (t.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
             .replace(_ESC_STAR, '*').replace(_ESC_DOLLAR, '$'))


def _run(text, bold=False, italic=False, vert=None, east_asia=None):
    rpr = '<w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:cs="Times New Roman"'
    if east_asia:
        rpr += f' w:eastAsia="{east_asia}"'
    rpr += '/>'
    if bold:
        rpr += '<w:b/>'
    if italic:
        rpr += '<w:i/>'
    if vert:
        rpr += f'<w:vertAlign w:val="{vert}"/>'
    if east_asia:
        rpr += '<w:lang w:val="en-US" w:eastAsia="ja-JP"/>'
    return (f'<w:r><w:rPr>{rpr}</w:rPr>'
            f'<w:t xml:space="preserve">{_esc(text)}</w:t></w:r>')


_MATH_ITALIC = re.compile(r'[A-Za-zα-ωΑ-Ω]')


_UPRIGHT_WORDS = {'min', 'max', 'log', 'ops', 'P50'}


def _math_runs(expr, east_asia=None, bold=False):
    """$...$ の中身を runs に。_x/_{...} 下付き, ^x/^{...} 上付き, 英字イタリック。
    下付き/上付き内の 2 文字以上の英字列（opt, res, RSR, adj 等のラベル）と
    min/max 等の関数語は立体にする。"""
    out = []

    def emit(seg, vert=None):
        for tok in re.findall(r'[A-Za-z]+|[α-ωΑ-Ω]|.', seg, re.S):
            if tok.isalpha() and len(tok) > 1 and re.match(r'[A-Za-z]', tok):
                if vert is not None or tok in _UPRIGHT_WORDS:
                    out.append(_run(tok, bold=bold, vert=vert,
                                    east_asia=east_asia))
                else:
                    for c in tok:
                        out.append(_run(c, bold=bold, italic=True,
                                        east_asia=east_asia))
            else:
                out.append(_run(tok, bold=bold,
                                italic=bool(_MATH_ITALIC.match(tok)),
                                vert=vert, east_asia=east_asia))

    i = 0
    while i < len(expr):
        ch = expr[i]
        if ch in '_^':
            vert = 'subscript' if ch == '_' else 'superscript'
            i += 1
            if i < len(expr) and expr[i] == '{':
                j = expr.index('}', i)
                seg = expr[i + 1:j]
                i = j + 1
            else:
                seg = expr[i]
                i += 1
            emit(seg, vert)
            continue
        emit(ch)
        i += 1
    return ''.join(out)


def rich_runs(text, east_asia=None, base_bold=False, base_italic=False):
    """**bold** / *italic* / $math$（入れ子対応: 太字内の $math$ も処理）。"""
    text = _protect_escapes(text)
    parts = _TOKEN.split(text)
    out = []
    for p in parts:
        if not p:
            continue
        if p.startswith('**') and p.endswith('**') and len(p) > 4:
            out.append(rich_runs(p[2:-2], east_asia, base_bold=True,
                                 base_italic=base_italic))
        elif p.startswith('*') and p.endswith('*') and len(p) > 2:
            out.append(rich_runs(p[1:-1], east_asia, base_bold=base_bold,
                                 base_italic=True))
        elif p.startswith('$') and p.endswith('$'):
            # 数式 run は純 Latin として扱う（eastAsia 言語を継がせない）
            out.append(_math_runs(p[1:-1], east_asia=None, bold=base_bold))
        else:
            out.append(_run(p, bold=base_bold, italic=base_italic,
                            east_asia=east_asia))
    return ''.join(out)


# ---------------- 段落テンプレート ----------------

LINE_EXACT = '<w:spacing w:line="240" w:lineRule="exact"/>'


def p_xml(inner, ppr=''):
    return f'<w:p {NS_W} {NS_M}><w:pPr>{ppr}</w:pPr>{inner}</w:p>'


def para_title(text, ea):
    return p_xml(rich_runs(text, ea).replace('</w:rPr>',
                 '<w:bCs/><w:sz w:val="40"/><w:szCs w:val="40"/></w:rPr>'),
                 '<w:jc w:val="center"/>')


def para_center(text, ea, bold=False):
    runs = rich_runs(text, ea, base_bold=bold)
    return p_xml(runs, LINE_EXACT + '<w:jc w:val="center"/>')


def para_abstract(text, ea):
    lead = _run('Abstract. ', bold=True, east_asia=ea)
    body = rich_runs(text, ea)
    return p_xml(lead + body,
                 LINE_EXACT + '<w:ind w:left="680" w:right="680"/>'
                 '<w:jc w:val="both"/>')


def para_keywords(label, text, ea):
    return p_xml(_run(label + ' ', bold=True, east_asia=ea) + rich_runs(text, ea),
                 LINE_EXACT + '<w:ind w:left="680" w:right="680"/>')


def para_heading(text, ea):
    runs = rich_runs(text, ea).replace(
        '</w:rPr>', '<w:sz w:val="22"/><w:szCs w:val="22"/></w:rPr>')
    return p_xml(runs, '<w:pStyle w:val="2"/>'
                       '<w:spacing w:before="90" w:after="20"'
                       ' w:line="240" w:lineRule="exact"/>')


_LEAD_PREFIX = re.compile(r'^\s*(?:\d+\.\s+|•\s+|\(\w\)\s+)?')


def _apply_emphasis_policy(text):
    """テンプレ準拠: 強調はイタリック・控えめに。ただし段落頭の run-in ラベル
    （小見出し的な先頭太字）は太字のまま残し、文中の語句強調 **...** はイタリック
    *...* に落とす。番号 "1. " や中黒 "•" 始まりのラベルも先頭扱い。"""
    spans = [(m.start(), m.end()) for m in re.finditer(r'\*\*.+?\*\*', text)]
    if not spans:
        return text
    lead_end = _LEAD_PREFIX.match(text).end()
    keep_first = spans[0][0] == lead_end        # 先頭ラベルなら太字維持
    out, last = [], 0
    for idx, (s, e) in enumerate(spans):
        out.append(text[last:s])
        inner = text[s + 2:e - 2]
        out.append(f'**{inner}**' if (idx == 0 and keep_first) else f'*{inner}*')
        last = e
    out.append(text[last:])
    return ''.join(out)


def para_body(text, ea, first_indent=True, demote_bold=False):
    if demote_bold:
        text = _apply_emphasis_policy(text)
    ind = '' if first_indent else '<w:ind w:firstLineChars="0" w:firstLine="0"/>'
    return p_xml(rich_runs(text, ea),
                 f'<w:pStyle w:val="a8"/>{LINE_EXACT}{ind}<w:jc w:val="both"/>')


def para_empty(ea=None, tiny=False, half=False):
    sz = '<w:rPr><w:sz w:val="2"/></w:rPr>' if tiny else ''
    if tiny:
        sp = '<w:spacing w:line="14" w:lineRule="exact"/>'
    elif half:
        sp = '<w:spacing w:line="120" w:lineRule="exact"/>'
        sz = '<w:rPr><w:sz w:val="12"/></w:rPr>'
    else:
        sp = LINE_EXACT
    return f'<w:p {NS_W}><w:pPr>{sp}{sz}</w:pPr></w:p>'


def para_sect_break(cols):
    """直前までの内容を cols 段のセクションとして閉じる極小段落。"""
    return (f'<w:p {NS_W}><w:pPr>'
            f'<w:spacing w:line="14" w:lineRule="exact"/>'
            f'{sect_xml(cols)}'
            f'<w:rPr><w:sz w:val="2"/></w:rPr></w:pPr></w:p>')


def para_caption(text, ea):
    return p_xml(rich_runs(text, ea),
                 LINE_EXACT + '<w:jc w:val="center"/>'
                 '<w:ind w:left="284" w:right="284"/>')


def para_equation(omml, number):
    """中央タブ+右タブで  [式(中央)]  (n)[右] を1行に置く。"""
    tabs = (f'<w:tabs><w:tab w:val="center" w:pos="{COL_W // 2}"/>'
            f'<w:tab w:val="right" w:pos="{COL_W}"/></w:tabs>')
    inner = ('<w:r><w:tab/></w:r>' + omml +
             '<w:r><w:tab/></w:r>' + _run(number))
    # テンプレ指示: 数式の上下に1行ぶんの余白を空ける
    spacing = '<w:spacing w:before="120" w:after="120" w:line="330" w:lineRule="exact"/>'
    return p_xml(inner, spacing + tabs)


def para_ref(text, ea):
    runs = rich_runs(text, ea).replace(
        '</w:rPr>', '<w:sz w:val="17"/><w:szCs w:val="17"/></w:rPr>')
    return p_xml(runs,
                 '<w:spacing w:line="190" w:lineRule="exact"/>'
                 '<w:ind w:left="284" w:hanging="284"/><w:jc w:val="both"/>')


# ---------------- テーブル ----------------

def table_xml(header, rows, widths, ea, total_w=COL_W, fs=16):
    """widths: 比率リスト。fs: half-points（16=8pt）。"""
    tw = [int(total_w * w / sum(widths)) for w in widths]

    def cell(text, w, bold=False, top=False, bottom=False):
        borders = '<w:tcBorders>'
        if top:
            borders += '<w:top w:val="single" w:sz="8" w:space="0" w:color="000000"/>'
        if bottom:
            borders += '<w:bottom w:val="single" w:sz="8" w:space="0" w:color="000000"/>'
        borders += '</w:tcBorders>'
        runs = rich_runs(text, ea, base_bold=bold).replace(
            '</w:rPr>', f'<w:sz w:val="{fs}"/><w:szCs w:val="{fs}"/></w:rPr>')
        return (f'<w:tc><w:tcPr><w:tcW w:w="{w}" w:type="dxa"/>{borders}</w:tcPr>'
                f'<w:p><w:pPr><w:spacing w:line="180" w:lineRule="exact"/></w:pPr>'
                f'{runs}</w:p></w:tc>')

    xml = (f'<w:tbl {NS_W}><w:tblPr><w:tblW w:w="{total_w}" w:type="dxa"/>'
           '<w:tblLayout w:type="fixed"/></w:tblPr><w:tblGrid>' +
           ''.join(f'<w:gridCol w:w="{w}"/>' for w in tw) + '</w:tblGrid>')
    xml += '<w:tr>' + ''.join(
        cell(h, tw[j], bold=True, top=True, bottom=True)
        for j, h in enumerate(header)) + '</w:tr>'
    for i, row in enumerate(rows):
        last = i == len(rows) - 1
        xml += '<w:tr>' + ''.join(
            cell(c, tw[j], bottom=last) for j, c in enumerate(row)) + '</w:tr>'
    xml += '</w:tbl>'
    return xml


# ---------------- ビルド本体 ----------------

def _first_sectpr_copy(doc):
    """テンプレ先頭セクション(1段+titlePg+ヘッダ参照)の sectPr を複製して返す。"""
    body = doc.element.body
    for p in body.iter(qn('w:p')):
        pPr = p.find(qn('w:pPr'))
        if pPr is not None and pPr.find(qn('w:sectPr')) is not None:
            return copy.deepcopy(pPr.find(qn('w:sectPr')))
    raise RuntimeError('template first sectPr not found')


def build(template_path, out_path, blocks, east_asia=None, fig_dir='',
          demote_bold=False):
    doc = Document(template_path)
    body = doc.element.body
    title_sect = _first_sectpr_copy(doc)
    # ヘッダ・フッタはテンプレ流用物（2016年編者名・ページ番号）なので全て外す
    for tag in ('w:headerReference', 'w:footerReference'):
        for el in title_sect.findall(qn(tag)):
            title_sect.remove(el)
    # テンプレ styles の wordWrap=0（英単語の途中改行を許す）を無効化
    for ww in doc.styles.element.iter(qn('w:wordWrap')):
        ww.set(qn('w:val'), '1')
    # 自動ハイフネーション（両端揃え 2 段組の行送り効率を上げる）。
    # CT_Settings は要素順が厳密なので defaultTabStop の直後に入れる。
    st = doc.settings.element
    if st.find(qn('w:autoHyphenation')) is None:
        tab = st.find(qn('w:defaultTabStop'))
        el = parse_xml(f'<w:autoHyphenation {NS_W}/>')
        if tab is not None:
            tab.addnext(el)
        else:
            st.insert(0, el)

    # 既存本文を全削除（末尾 sectPr は残す）
    for el in list(body):
        if el.tag != qn('w:sectPr'):
            body.remove(el)
    final_sect = body.find(qn('w:sectPr'))
    # 最終セクションを 2 段 continuous に強制
    for el in list(final_sect):
        final_sect.remove(el)
    tmp = parse_xml(sect_xml(2))
    for el in list(tmp):
        final_sect.append(copy.deepcopy(el))

    def add(xml_str):
        el = parse_xml(xml_str)
        body.insert(len(body) - 1, el)  # final sectPr の手前
        return el

    ea = east_asia
    pending_title_sect = False
    after_heading = False

    for blk in blocks:
        kind = blk[0]
        if kind == 'title':
            add(para_title(blk[1], ea))
        elif kind == 'authors':
            add(para_empty(ea, half=True))
            for name in blk[1]:
                add(para_center(name, ea, bold=True))
        elif kind == 'affil':
            for line in blk[1]:
                add(para_center(line, ea))
        elif kind == 'abstract':
            add(para_empty(ea, half=True))
            add(para_abstract(blk[1], ea))
        elif kind == 'keywords':
            add(para_empty(ea, half=True))
            add(para_keywords(blk[1], blk[2], ea))
            add(para_empty(ea))
            # タイトル部セクション（1段, titlePg, ヘッダ付き）をここで閉じる
            p = add(f'<w:p {NS_W}><w:pPr><w:spacing w:line="14" w:lineRule="exact"/>'
                    f'<w:rPr><w:sz w:val="2"/></w:rPr></w:pPr></w:p>')
            p.find(qn('w:pPr')).append(title_sect)
            pending_title_sect = True
        elif kind in ('h1', 'h2', 'h3'):
            add(para_heading(blk[1], ea))
            after_heading = True
            continue
        elif kind == 'p':
            add(para_body(blk[1], ea, first_indent=True, demote_bold=demote_bold))
        elif kind == 'p_noindent':
            add(para_body(blk[1], ea, first_indent=False, demote_bold=demote_bold))
        elif kind == 'eq':
            add(para_equation(blk[1], blk[2]))
        elif kind == 'fig':
            path, caption, mode = blk[1], blk[2], blk[3]
            scale = blk[4] if len(blk) > 4 else 1.0
            from PIL import Image
            img = Image.open(os.path.join(fig_dir, path))
            w_px, h_px = img.size
            box_w = TEXT_W if mode == 'full' else COL_W
            w_emu = int(Twips(int(box_w * scale)))
            h_emu = int(w_emu * h_px / w_px)
            if mode == 'full':
                # ページ上端フロートの 1 列テーブル（画像 + キャプション）。
                # 2 段組の上に全幅で載り、本文は下に回り込む（空白が出ない）。
                _add_float_figure(doc, body, os.path.join(fig_dir, path),
                                  w_emu, h_emu, caption, ea)
            else:
                pic_para = doc.add_paragraph()
                ppr = pic_para._p.get_or_add_pPr()
                ppr.append(parse_xml(
                    f'<w:spacing {NS_W} w:line="240" w:lineRule="auto"/>'))
                ppr.append(parse_xml(f'<w:jc {NS_W} w:val="center"/>'))
                run = pic_para.add_run()
                run.add_picture(os.path.join(fig_dir, path),
                                width=Emu(w_emu), height=Emu(h_emu))
                add(para_caption(caption, ea))
        elif kind == 'table':
            add(para_caption(blk[4], ea))  # 表キャプションは上
            add(table_xml(blk[1], blk[2], blk[3], ea))
            add(para_empty(ea))
        elif kind == 'ref':
            add(para_ref(blk[1], ea))
        elif kind == 'raw':
            add(blk[1])
        else:
            raise ValueError(f'unknown block: {kind}')
        after_heading = False

    if not pending_title_sect:
        raise RuntimeError('keywords block (title section close) missing')
    doc.save(out_path)
    print(' ->', out_path)


def _add_float_figure(doc, body, img_path, w_emu, h_emu, caption, ea):
    """ページ上端フロートの 1 列×2 行テーブル（画像/キャプション）を挿入する。"""
    cap_runs = rich_runs(caption, ea)
    tbl_xml = (
        f'<w:tbl {NS_W} {NS_M}><w:tblPr>'
        f'<w:tblW w:w="{TEXT_W}" w:type="dxa"/>'
        '<w:tblpPr w:leftFromText="141" w:rightFromText="141"'
        ' w:topFromText="141" w:bottomFromText="200"'
        ' w:vertAnchor="margin" w:horzAnchor="margin"'
        ' w:tblpXSpec="center" w:tblpYSpec="top"/>'
        '<w:tblOverlap w:val="never"/>'
        '<w:tblLayout w:type="fixed"/>'
        '<w:tblCellMar><w:left w:w="0" w:type="dxa"/>'
        '<w:right w:w="0" w:type="dxa"/></w:tblCellMar>'
        f'</w:tblPr><w:tblGrid><w:gridCol w:w="{TEXT_W}"/></w:tblGrid>'
        f'<w:tr><w:tc><w:tcPr><w:tcW w:w="{TEXT_W}" w:type="dxa"/></w:tcPr>'
        '<w:p><w:pPr><w:spacing w:after="60" w:line="240" w:lineRule="auto"/>'
        '<w:jc w:val="center"/></w:pPr></w:p></w:tc></w:tr>'
        f'<w:tr><w:tc><w:tcPr><w:tcW w:w="{TEXT_W}" w:type="dxa"/></w:tcPr>'
        f'<w:p><w:pPr>{LINE_EXACT}<w:jc w:val="center"/>'
        f'<w:ind w:left="284" w:right="284"/></w:pPr>{cap_runs}</w:p>'
        '</w:tc></w:tr></w:tbl>')
    tbl = parse_xml(tbl_xml)
    body.insert(len(body) - 1, tbl)
    # 画像を 1 行目セルの段落へ
    from docx.text.paragraph import Paragraph
    pic_p = tbl.find(qn('w:tr')).find(qn('w:tc')).find(qn('w:p'))
    para = Paragraph(pic_p, doc._body)
    run = para.add_run()
    run.add_picture(img_path, width=Emu(w_emu), height=Emu(h_emu))
    # フロート表の直後に極小の通常段落を置く（表が連続すると結合されるため）
    body.insert(len(body) - 1, parse_xml(
        f'<w:p {NS_W}><w:pPr><w:spacing w:line="14" w:lineRule="exact"/>'
        f'<w:rPr><w:sz w:val="2"/></w:rPr></w:pPr></w:p>'))


# ---------------- OMML 数式 ----------------

def _mr(t, sty=None):
    rpr = ('<w:rPr><w:rFonts w:ascii="Cambria Math" w:hAnsi="Cambria Math"/></w:rPr>')
    sty_x = f'<m:rPr><m:sty m:val="{sty}"/></m:rPr>' if sty else ''
    return f'<m:r>{sty_x}{rpr}<m:t xml:space="preserve">{_esc(t)}</m:t></m:r>'


def _ssub(base, sub):
    return (f'<m:sSub><m:e>{base}</m:e><m:sub>{sub}</m:sub></m:sSub>')


def _ssubsup(base, sub, sup):
    return (f'<m:sSubSup><m:e>{base}</m:e><m:sub>{sub}</m:sub>'
            f'<m:sup>{sup}</m:sup></m:sSubSup>')


def _hat(base):
    return ('<m:acc><m:accPr><m:chr m:val="̂"/></m:accPr>'
            f'<m:e>{base}</m:e></m:acc>')


def omml_eq1():
    """D(S_p,S_q) = Σ_{(i,j)∈O_opt} |r_{i,j}^p − r_{i,j}^q|"""
    Sp = _ssub(_mr('S'), _mr('p'))
    Sq = _ssub(_mr('S'), _mr('q'))
    sub = (_mr('(') + _mr('i') + _mr(',') + _mr('j') + _mr(')') + _mr('∈') +
           _ssub(_mr('O'), _mr('opt')))
    term = (_ssubsup(_mr('r'), _mr('i') + _mr(',') + _mr('j'), _mr('p')) +
            _mr('−') +
            _ssubsup(_mr('r'), _mr('i') + _mr(',') + _mr('j'), _mr('q')))
    absd = ('<m:d><m:dPr><m:begChr m:val="|"/><m:endChr m:val="|"/></m:dPr>'
            f'<m:e>{term}</m:e></m:d>')
    nary = ('<m:nary><m:naryPr><m:chr m:val="∑"/><m:limLoc m:val="undOvr"/>'
            '<m:supHide m:val="1"/></m:naryPr>'
            f'<m:sub>{sub}</m:sub><m:sup/><m:e>{absd}</m:e></m:nary>')
    return (f'<m:oMath {NS_M} {NS_W}>' +
            _mr('D') + _mr('(') + Sp + _mr(',') + Sq + _mr(')') + _mr('=') +
            nary + '</m:oMath>')


def omml_eq2():
    """F(S_q) = λ D̂(S_p,S_q) + (1−λ) M̂S(S_q)"""
    Sp = _ssub(_mr('S'), _mr('p'))
    Sq = _ssub(_mr('S'), _mr('q'))

    def Sq2():
        return _ssub(_mr('S'), _mr('q'))

    return (f'<m:oMath {NS_M} {NS_W}>' +
            _mr('F') + _mr('(') + Sq + _mr(')') + _mr('=') + _mr('λ') +
            _hat(_mr('D')) + _mr('(') + Sp + _mr(',') + Sq2() + _mr(')') +
            _mr('+') + _mr('(') + _mr('1') + _mr('−') + _mr('λ') +
            _mr(')') + _hat(_mr('MS')) + _mr('(') + Sq2() + _mr(')') +
            '</m:oMath>')
